"""
Full organization report generator — DOCX with all data: objects, apartments,
clients, payments, expenses, suppliers, debts. Saved to OrganizationReport and MEDIA.
"""
import os
from decimal import Decimal
from django.conf import settings
from django.utils import timezone
from django.db.models import Sum
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import (
    Object, Apartment, User, Payment, Expense, Supplier,
    ExpenseType, UserPayment, SupplierPayment, OrganizationReport,
)


def _add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def _add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


def _add_table_from_rows(doc, headers, rows):
    """Add a table with headers and data rows. headers and rows are lists of strings/cells."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hrow = table.rows[0].cells
    for i, h in enumerate(headers):
        hrow[i].text = str(h)
    for ri, row in enumerate(rows):
        r = table.rows[ri + 1].cells
        for i, cell in enumerate(row):
            if i < len(r):
                r[i].text = str(cell) if cell is not None else ''
    return table


def build_full_report_docx():
    """
    Builds a full organization report DOCX: Ahlan House hisoboti + current datetime,
    then objects, apartments, clients, payments (kimga qancha), user payments (kimdan qancha),
    expenses, suppliers, supplier payments, qarzdorlar. Saves file to MEDIA/reports/
    and creates OrganizationReport. Returns the OrganizationReport instance.
    """
    now = timezone.now()
    title = f"Ahlan House hisoboti — {now.strftime('%Y-%m-%d %H:%M')}"
    doc = DocxDocument()

    # Title and date
    _add_heading(doc, "Ahlan House hisoboti", level=0)
    _add_para(doc, f"Hisobot sanasi va vaqti: {now.strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph()

    # --- 1. Obyektlar ---
    _add_heading(doc, "1. Obyektlar", level=1)
    objects = Object.objects.all().order_by('id')
    if objects:
        _add_table_from_rows(doc,
            ['№', 'Nomi', 'Xonadonlar soni', 'Qavatlar', 'Manzil'],
            [
                (i, o.name, o.total_apartments, o.floors, (o.address or '')[:60])
                for i, o in enumerate(objects, 1)
            ]
        )
    else:
        _add_para(doc, "Obyektlar yo'q.")
    doc.add_paragraph()

    # --- 2. Xonadonlar ---
    _add_heading(doc, "2. Xonadonlar", level=1)
    apartments = Apartment.objects.select_related('object').all().order_by('object__name', 'room_number')
    if apartments:
        _add_table_from_rows(doc,
            ['№', 'Obyekt', 'Xona raqami', 'Xonalar', 'Maydoni', 'Narx', 'Holat', 'Balans', 'Jami to\'lov'],
            [
                (i, a.object.name, a.room_number, a.rooms, a.area, a.price, a.get_status_display(), a.balance, a.total_payments)
                for i, a in enumerate(apartments, 1)
            ]
        )
    else:
        _add_para(doc, "Xonadonlar yo'q.")
    doc.add_paragraph()

    # --- 3. Mijozlar (users) ---
    _add_heading(doc, "3. Mijozlar", level=1)
    clients = User.objects.filter(user_type='mijoz').order_by('id')
    if clients:
        _add_table_from_rows(doc,
            ['№', 'F.I.O', 'Telefon', 'Manzil', 'Balans'],
            [(i, c.fio, c.phone_number or '', (c.address or '')[:40], c.balance) for i, c in enumerate(clients, 1)]
        )
    else:
        _add_para(doc, "Mijozlar yo'q.")
    doc.add_paragraph()

    # --- 4. To'lovlar (kimga qancha — xonadon bo'yicha shartnoma to'lovlari) ---
    _add_heading(doc, "4. To'lovlar (shartnoma bo'yicha — kimga qancha berilgan)", level=1)
    payments = Payment.objects.select_related('user', 'apartment', 'apartment__object').all().order_by('-created_at')
    if payments:
        _add_table_from_rows(doc,
            ['№', 'Mijoz', 'Xonadon', 'Jami summa', 'To\'langan', 'Qolgan', 'Holat', 'To\'lov turi'],
            [
                (i, p.user.fio, f"{p.apartment.object.name} — {p.apartment.room_number}", p.total_amount, p.paid_amount, p.total_amount - p.paid_amount, p.get_status_display(), p.get_payment_type_display())
                for i, p in enumerate(payments, 1)
            ]
        )
    else:
        _add_para(doc, "To'lovlar yo'q.")
    doc.add_paragraph()

    # --- 5. Mijozdan olingan to'lovlar (UserPayment) ---
    _add_heading(doc, "5. Mijozlardan olingan to'lovlar (balansga kiritilgan)", level=1)
    user_payments = UserPayment.objects.select_related('user').all().order_by('-date')
    if user_payments:
        _add_table_from_rows(doc,
            ['№', 'Mijoz', 'Summa', 'To\'lov turi', 'Sana'],
            [(i, up.user.fio, up.amount, up.get_payment_type_display(), up.date.strftime('%d.%m.%Y %H:%M')) for i, up in enumerate(user_payments, 1)]
        )
    else:
        _add_para(doc, "Mijozlardan olingan to'lovlar yo'q.")
    doc.add_paragraph()

    # --- 6. Xarajatlar ---
    _add_heading(doc, "6. Xarajatlar", level=1)
    expenses = Expense.objects.select_related('supplier', 'expense_type', 'object').all().order_by('-date')
    if expenses:
        _add_table_from_rows(doc,
            ['№', 'Yetkazib beruvchi', 'Xarajat turi', 'Obyekt', 'Summa', 'Sana', 'Izoh', 'Holat'],
            [(i, e.supplier.company_name, e.expense_type.name, e.object.name, e.amount, e.date, (e.comment or '')[:30], e.status) for i, e in enumerate(expenses, 1)]
        )
    else:
        _add_para(doc, "Xarajatlar yo'q.")
    doc.add_paragraph()

    # --- 7. Yetkazib beruvchilar va balans ---
    _add_heading(doc, "7. Yetkazib beruvchilar (balans — bizdan qarz)", level=1)
    suppliers = Supplier.objects.all().order_by('company_name')
    if suppliers:
        _add_table_from_rows(doc,
            ['№', 'Kompaniya', 'Telefon', 'Balans (bizdan qarz)'],
            [(i, s.company_name, s.phone_number or '', s.balance) for i, s in enumerate(suppliers, 1)]
        )
    else:
        _add_para(doc, "Yetkazib beruvchilar yo'q.")
    doc.add_paragraph()

    # --- 8. Yetkazib beruvchilarga berilgan to'lovlar ---
    _add_heading(doc, "8. Yetkazib beruvchilarga berilgan to'lovlar", level=1)
    supplier_payments = SupplierPayment.objects.select_related('supplier').all().order_by('-date')
    if supplier_payments:
        _add_table_from_rows(doc,
            ['№', 'Yetkazib beruvchi', 'Summa', 'To\'lov turi', 'Sana'],
            [(i, sp.supplier.company_name, sp.amount, sp.get_payment_type_display(), sp.date.strftime('%d.%m.%Y %H:%M')) for i, sp in enumerate(supplier_payments, 1)]
        )
    else:
        _add_para(doc, "Yetkazib beruvchilarga to'lovlar yo'q.")
    doc.add_paragraph()

    # --- 9. Qarzdorlar (muddati o'tgan to'lovlar) ---
    _add_heading(doc, "9. Qarzdorlar (muddati o'tgan to'lovlar)", level=1)
    overdue_list = []
    for payment in Payment.objects.filter(
        payment_type='muddatli', status__in=['pending', 'overdue']
    ).select_related('user', 'apartment', 'apartment__object'):
        data = payment.get_overdue_payments()
        total_overdue = data.get('total_overdue') or Decimal('0')
        if total_overdue > 0:
            overdue_list.append((
                payment.apartment.object.name,
                payment.apartment.room_number,
                payment.user.fio,
                total_overdue,
                len(data.get('overdue_payments', []))
            ))
    if overdue_list:
        _add_table_from_rows(doc,
            ['Obyekt', 'Xonadon', 'Mijoz', 'Qarz summa', 'Kechikkan oylar'],
            overdue_list
        )
    else:
        _add_para(doc, "Qarzdorlar yo'q.")
    doc.add_paragraph()

    # --- 10. Umumiy qisqacha ---
    _add_heading(doc, "10. Umumiy qisqacha", level=1)
    total_objects = Object.objects.count()
    total_apartments = Apartment.objects.count()
    sold = Apartment.objects.filter(status='sotilgan').count()
    free = Apartment.objects.filter(status='bosh').count()
    total_sales = Payment.objects.aggregate(s=Sum('total_amount'))['s'] or Decimal('0')
    total_paid = Payment.objects.aggregate(s=Sum('paid_amount'))['s'] or Decimal('0')
    total_expenses_sum = Expense.objects.aggregate(s=Sum('amount'))['s'] or Decimal('0')
    _add_para(doc, f"Obyektlar: {total_objects}, Xonadonlar: {total_apartments}, Sotilgan: {sold}, Bo'sh: {free}.")
    _add_para(doc, f"Shartnoma bo'yicha jami summa: {total_sales}, To'langan: {total_paid}, Qolgan: {total_sales - total_paid}.")
    _add_para(doc, f"Xarajatlar jami: {total_expenses_sum}.")

    # Save file and create OrganizationReport
    os.makedirs(os.path.join(settings.MEDIA_ROOT, 'reports'), exist_ok=True)
    safe_ts = now.strftime('%Y-%m-%d_%H-%M')
    filename = f"ahlan_house_hisoboti_{safe_ts}.docx"
    file_path = os.path.join(settings.MEDIA_ROOT, 'reports', filename)
    doc.save(file_path)

    report = OrganizationReport.objects.create(
        title=title,
        file=f'reports/{filename}'
    )
    return report
