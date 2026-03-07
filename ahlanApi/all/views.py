from rest_framework import viewsets, permissions, status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.parsers import JSONParser, MultiPartParser, FormParser
from .models import Object, Apartment, User, ExpenseType, Supplier, Expense, Payment, Document, UserPayment, SupplierPayment, RoomTypeModel, OrganizationReport
from .serializers import (
    ObjectSerializer, ApartmentSerializer, UserSerializer,
    ExpenseTypeSerializer, SupplierSerializer, ExpenseSerializer,
    PaymentSerializer, UserPaymentSerializer, DocumentSerializer, SupplierPaymentSerializer,
    RoomTypeModelSerializer, OrganizationReportSerializer,
)
from .report_service import build_full_report_docx
from .pagination import CustomPagination
from .filters import PaymentFilter
from rest_framework import serializers
from rest_framework_simplejwt.views import TokenObtainPairView
from rest_framework_simplejwt.serializers import TokenObtainPairSerializer
from django.http import FileResponse
from django.core.files import File
import io
import os
import zipfile
import shutil
import tempfile
from django.conf import settings
from django.db.models import Sum, Count, Avg
from datetime import datetime
from decimal import Decimal
from docx import Document as DocxDocument
from django.utils import timezone


def _convert_gltf_to_metalrough(file_path_abs):
    """
    Convert glTF/GLB from KHR_materials_pbrSpecularGlossiness to metallic-roughness
    so model-viewer/Three.js can load it. Uses npx @gltf-transform/cli metalrough.
    On failure (no Node, or not needed), leaves file unchanged.
    """
    if not file_path_abs or not os.path.isfile(file_path_abs):
        return
    low = file_path_abs.lower()
    if not (low.endswith('.glb') or low.endswith('.gltf')):
        return
    import subprocess
    dir_name = os.path.dirname(file_path_abs)
    base = os.path.basename(file_path_abs)
    out_path = os.path.join(dir_name, '_converted_' + base)
    try:
        if os.name == 'nt':
            subprocess.run(
                ['cmd', '/c', 'npx', '--yes', '@gltf-transform/cli', 'metalrough', file_path_abs, out_path],
                capture_output=True,
                timeout=120,
                cwd=dir_name,
            )
        else:
            subprocess.run(
                ['npx', '--yes', '@gltf-transform/cli', 'metalrough', file_path_abs, out_path],
                capture_output=True,
                timeout=120,
                cwd=dir_name,
            )
        if os.path.isfile(out_path):
            os.replace(out_path, file_path_abs)
    except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
        if os.path.isfile(out_path):
            try:
                os.remove(out_path)
            except OSError:
                pass


def _extract_zip_and_find_model(zip_file, target_dir_abs):
    """
    Extract zip into target_dir_abs (absolute path under MEDIA_ROOT).
    Returns relative path (from MEDIA_ROOT) to first .glb or .gltf, or None.
    Safe: no path traversal outside target_dir.
    """
    os.makedirs(target_dir_abs, exist_ok=True)
    media_root = os.path.realpath(settings.MEDIA_ROOT)
    target_real = os.path.realpath(target_dir_abs)
    if not target_real.startswith(media_root):
        return None
    main_relative = None
    try:
        with zipfile.ZipFile(zip_file, 'r') as z:
            for name in z.namelist():
                if name.endswith('/'):
                    continue
                name_safe = os.path.normpath(name)
                if '..' in name_safe or name_safe.startswith('/'):
                    continue
                dest = os.path.join(target_dir_abs, name_safe)
                dest_real = os.path.realpath(dest)
                if not dest_real.startswith(target_real):
                    continue
                z.extract(name, target_dir_abs)
                if main_relative is None and (name_safe.lower().endswith('.glb') or name_safe.lower().endswith('.gltf')):
                    main_relative = os.path.join(os.path.relpath(target_dir_abs, settings.MEDIA_ROOT), name_safe).replace('\\', '/')
        return main_relative
    except (zipfile.BadZipFile, OSError):
        return None

class CustomTokenObtainPairSerializer(TokenObtainPairSerializer):
    """Accepts phone_number + password from frontend; maps phone_number to username for Django auth."""
    phone_number = serializers.CharField(required=False)

    def validate(self, attrs):
        # Frontend sends phone_number; SimpleJWT uses username (our User.USERNAME_FIELD = phone_number)
        phone = attrs.get('phone_number') or attrs.get('username') or ''
        if phone and not attrs.get('username'):
            phone = str(phone).strip()
            if phone and not phone.startswith('+'):
                phone = '+' + phone
            attrs['username'] = phone
        return super().validate(attrs)

    @classmethod
    def get_token(cls, user):
        token = super().get_token(user)
        token['user_type'] = user.user_type
        token['fio'] = user.fio
        return token


class CustomTokenObtainPairView(TokenObtainPairView):
    serializer_class = CustomTokenObtainPairSerializer

class ObjectViewSet(viewsets.ModelViewSet):
    queryset = Object.objects.all()
    serializer_class = ObjectSerializer
    pagination_class = CustomPagination
    permission_classes = [permissions.IsAuthenticated]
    filterset_fields = ['name', 'floors', 'total_apartments']

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

    @action(detail=True, methods=['post'], permission_classes=[permissions.IsAuthenticated])
    def upload_3d(self, request, pk=None):
        """Obyekt uchun asosiy 3D model. .glb/.gltf yoki .zip (ichida model + teksturalar). multipart/form-data, file: model_3d"""
        obj = self.get_object()
        file = request.FILES.get('model_3d')
        if not file:
            return Response({'detail': 'model_3d fayl yuborilishi shart'}, status=status.HTTP_400_BAD_REQUEST)
        is_zip = file.name.lower().endswith('.zip')
        allowed = ('.glb', '.gltf', '.zip')
        if not is_zip and not any(file.name.lower().endswith(ext) for ext in ('.glb', '.gltf')):
            return Response({'detail': 'Faqat .glb, .gltf yoki .zip qabul qilinadi'}, status=status.HTTP_400_BAD_REQUEST)
        target_dir = os.path.join(settings.MEDIA_ROOT, f'object_3d/object_{obj.id}')
        if obj.model_3d:
            obj.model_3d.delete(save=False)
        if os.path.isdir(target_dir):
            shutil.rmtree(target_dir, ignore_errors=True)
        if is_zip:
            main_relative = _extract_zip_and_find_model(file, target_dir)
            if not main_relative:
                return Response({'detail': 'ZIP ichida .glb yoki .gltf topilmadi yoki xatolik'}, status=status.HTTP_400_BAD_REQUEST)
            prefix = f'object_3d/object_{obj.id}/'
            suffix = main_relative.split(prefix, 1)[-1] if prefix in main_relative else os.path.basename(main_relative)
            main_path = os.path.join(settings.MEDIA_ROOT, main_relative)
            with open(main_path, 'rb') as fh:
                obj.model_3d.save(suffix, File(fh), save=True)
        else:
            obj.model_3d = file
            obj.save()
        try:
            if obj.model_3d and obj.model_3d.name:
                full_path = os.path.join(settings.MEDIA_ROOT, obj.model_3d.name)
                if os.path.isfile(full_path):
                    _convert_gltf_to_metalrough(full_path)
        except Exception:
            pass
        serializer = self.get_serializer(obj)
        return Response(serializer.data, status=status.HTTP_200_OK)

    @action(detail=True, methods=['post'], permission_classes=[permissions.IsAuthenticated])
    def delete_3d(self, request, pk=None):
        """Obyektning asosiy 3D modelini o'chirish."""
        obj = self.get_object()
        if obj.model_3d:
            obj.model_3d.delete(save=False)
            obj.model_3d = None
            obj.save()
        target_dir = os.path.join(settings.MEDIA_ROOT, f'object_3d/object_{obj.id}')
        if os.path.isdir(target_dir):
            shutil.rmtree(target_dir, ignore_errors=True)
        serializer = self.get_serializer(obj)
        return Response(serializer.data, status=status.HTTP_200_OK)

    @action(detail=True, methods=['get'], permission_classes=[permissions.IsAuthenticated])
    def segments(self, request, pk=None):
        """Obyekt uchun segment ID lar ro'yxati (qavat_hujra). Asosiy modeldan ajratiladigan xonadon segmentlari."""
        obj = self.get_object()
        floors = max(1, obj.floors or 1)
        total = max(1, obj.total_apartments or 1)
        per_floor = max(1, (total + floors - 1) // floors)
        segments = []
        for f in range(1, floors + 1):
            for c in range(1, per_floor + 1):
                segments.append(f"{f}_{c}")
                if len(segments) >= total:
                    break
            if len(segments) >= total:
                break
        return Response({'segments': segments, 'floors': floors, 'per_floor': per_floor})

class ApartmentViewSet(viewsets.ModelViewSet):
    queryset = Apartment.objects.all()
    serializer_class = ApartmentSerializer
    pagination_class = CustomPagination
    permission_classes = [permissions.IsAuthenticated]
    filterset_fields = [
        'object', 'rooms', 'floor', 'status', 'price', 'area', 'room_number'
    ]

    def get_queryset(self):
        queryset = super().get_queryset()
        min_price = self.request.query_params.get('min_price', None)
        max_price = self.request.query_params.get('max_price', None)
        min_area = self.request.query_params.get('min_area', None)
        max_area = self.request.query_params.get('max_area', None)

        if min_price:
            queryset = queryset.filter(price__gte=min_price)
        if max_price:
            queryset = queryset.filter(price__lte=max_price)
        if min_area:
            queryset = queryset.filter(area__gte=min_area)
        if max_area:
            queryset = queryset.filter(area__lte=max_area)

        return queryset

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

    @action(detail=True, methods=['post'], permission_classes=[permissions.IsAuthenticated])
    def upload_3d(self, request, pk=None):
        """3D model: .glb/.gltf yoki .zip (ichida model + teksturalar). multipart/form-data, file: model_3d"""
        apartment = self.get_object()
        file = request.FILES.get('model_3d')
        if not file:
            return Response({'detail': 'model_3d fayl yuborilishi shart'}, status=status.HTTP_400_BAD_REQUEST)
        is_zip = file.name.lower().endswith('.zip')
        if not is_zip and not any(file.name.lower().endswith(ext) for ext in ('.glb', '.gltf')):
            return Response({'detail': 'Faqat .glb, .gltf yoki .zip qabul qilinadi'}, status=status.HTTP_400_BAD_REQUEST)
        target_dir = os.path.join(settings.MEDIA_ROOT, f'apartment_3d/apartment_{apartment.id}')
        if apartment.model_3d:
            apartment.model_3d.delete(save=False)
        if os.path.isdir(target_dir):
            shutil.rmtree(target_dir, ignore_errors=True)
        if is_zip:
            main_relative = _extract_zip_and_find_model(file, target_dir)
            if not main_relative:
                return Response({'detail': 'ZIP ichida .glb yoki .gltf topilmadi yoki xatolik'}, status=status.HTTP_400_BAD_REQUEST)
            prefix = f'apartment_3d/apartment_{apartment.id}/'
            suffix = main_relative.split(prefix, 1)[-1] if prefix in main_relative else os.path.basename(main_relative)
            main_path = os.path.join(settings.MEDIA_ROOT, main_relative)
            with open(main_path, 'rb') as fh:
                apartment.model_3d.save(suffix, File(fh), save=True)
        else:
            apartment.model_3d = file
            apartment.save()
        try:
            if apartment.model_3d and apartment.model_3d.name:
                full_path = os.path.join(settings.MEDIA_ROOT, apartment.model_3d.name)
                if os.path.isfile(full_path):
                    _convert_gltf_to_metalrough(full_path)
        except Exception:
            pass
        serializer = self.get_serializer(apartment)
        return Response(serializer.data, status=status.HTTP_200_OK)

    @action(detail=True, methods=['post'], permission_classes=[permissions.IsAuthenticated])
    def delete_3d(self, request, pk=None):
        """Xonadonning 3D model (override) ni o'chirish."""
        apartment = self.get_object()
        if apartment.model_3d:
            apartment.model_3d.delete(save=False)
            apartment.model_3d = None
            apartment.save()
        target_dir = os.path.join(settings.MEDIA_ROOT, f'apartment_3d/apartment_{apartment.id}')
        if os.path.isdir(target_dir):
            shutil.rmtree(target_dir, ignore_errors=True)
        serializer = self.get_serializer(apartment)
        return Response(serializer.data, status=status.HTTP_200_OK)

    @action(detail=True, methods=['post'], permission_classes=[permissions.IsAuthenticated])
    def set_segment(self, request, pk=None):
        """Xonadonni asosiy obyekt segmentiga biriktirish (segment_id). Body: {"segment_id": "1_1"} yoki null."""
        apartment = self.get_object()
        segment_id = request.data.get('segment_id')
        if segment_id is not None and segment_id != '' and not isinstance(segment_id, str):
            return Response({'detail': 'segment_id matn yoki bo\'sh bo\'lishi kerak'}, status=status.HTTP_400_BAD_REQUEST)
        apartment.segment_id = (segment_id or '').strip() or None
        apartment.save()
        serializer = self.get_serializer(apartment)
        return Response(serializer.data, status=status.HTTP_200_OK)

    @action(detail=True, methods=['get'], permission_classes=[permissions.IsAuthenticated])
    def get_total_payments(self, request, pk=None):
        apartment = self.get_object()
        total_payments = apartment.total_payments
        balance = apartment.balance
        return Response({
            'apartment': str(apartment),
            'total_payments': total_payments,
            'balance': balance
        })

    @action(detail=True, methods=['get'], permission_classes=[permissions.IsAuthenticated])
    def overdue_payments(self, request, pk=None):
        apartment = self.get_object()
        overdue_data = apartment.get_overdue_payments()
        return Response({
            'apartment': str(apartment),
            'overdue_payments': overdue_data['overdue_payments'],
            'total_overdue': overdue_data['total_overdue']
        })

    @action(detail=False, methods=['get'], permission_classes=[permissions.IsAuthenticated])
    def overdue_payments_report(self, request):
        object_id = request.query_params.get('object_id', None)
        apartment_id = request.query_params.get('apartment_id', None)
        queryset = Apartment.objects.all()

        if object_id:
            queryset = queryset.filter(object_id=object_id)
        if apartment_id:
            queryset = queryset.filter(id=apartment_id)

        report = []
        total_overdue_all = Decimal('0')
        for apartment in queryset:
            overdue_data = apartment.get_overdue_payments()
            if overdue_data['overdue_payments']:
                report.append({
                    'apartment': str(apartment),
                    'object': apartment.object.name,
                    'overdue_payments': overdue_data['overdue_payments'],
                    'total_overdue': overdue_data['total_overdue']
                })
                total_overdue_all += overdue_data['total_overdue']

        return Response({
            'report': report,
            'total_overdue_all': total_overdue_all
        })

class UserViewSet(viewsets.ModelViewSet):
    queryset = User.objects.all()
    serializer_class = UserSerializer
    pagination_class = CustomPagination
    permission_classes = [permissions.IsAuthenticated]
    filterset_fields = ['user_type', 'phone_number', 'balance']

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

    @action(detail=True, methods=['post'], permission_classes=[permissions.IsAdminUser])
    def add_balance(self, request, pk=None):
        user = self.get_object()
        amount = request.data.get('amount', 0)
        try:
            user.add_balance(float(amount))
            return Response({'message': f"{user.fio} balansiga {amount} so‘m qo‘shildi", 'balance': user.balance})
        except ValueError as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

class ExpenseTypeViewSet(viewsets.ModelViewSet):
    queryset = ExpenseType.objects.all()
    serializer_class = ExpenseTypeSerializer
    pagination_class = CustomPagination
    permission_classes = [permissions.IsAuthenticated]
    filterset_fields = ['name']

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

class SupplierViewSet(viewsets.ModelViewSet):
    queryset = Supplier.objects.all()
    serializer_class = SupplierSerializer
    pagination_class = CustomPagination
    permission_classes = [permissions.IsAuthenticated]
    filterset_fields = ['company_name', 'phone_number']

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

class ExpenseViewSet(viewsets.ModelViewSet):
    queryset = Expense.objects.all()
    serializer_class = ExpenseSerializer
    pagination_class = CustomPagination
    permission_classes = [permissions.IsAuthenticated]
    parser_classes = [JSONParser, MultiPartParser, FormParser]
    filterset_fields = ['date', 'supplier', 'expense_type', 'object', 'status']
    ordering_fields = ['date', 'id']  # Sana va ID bo'yicha tartiblashga ruxsat
    ordering = ['-date', '-id']

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

class PaymentViewSet(viewsets.ModelViewSet):
    queryset = Payment.objects.all().order_by("-created_at")
    serializer_class = PaymentSerializer
    pagination_class = CustomPagination
    permission_classes = [permissions.IsAuthenticated]
    filterset_class = PaymentFilter

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy', 'process_payment']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

    @action(detail=True, methods=['post'])
    def process_payment(self, request, pk=None):
        payment = self.get_object()
        try:
            amount = Decimal(str(request.data.get('amount', 0)))
        except (TypeError, ValueError):
            return Response({'error': 'Summa son bo\'lishi kerak'}, status=status.HTTP_400_BAD_REQUEST)
        payment_date = request.data.get('payment_date')

        if amount <= 0:
            return Response({'error': 'Summa musbat bo\'lishi kerak'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            if payment_date:
                parsed = datetime.strptime(
                    payment_date if isinstance(payment_date, str) else str(payment_date),
                    '%Y-%m-%d'
                )
                payment.payment_date = timezone.make_aware(parsed) if timezone.is_naive(parsed) else parsed
            payment.paid_amount += amount
            payment.update_status()
            payment.save()

            apartment = payment.apartment
            apartment.update_balance()
            apartment.update_status()

            serializer = self.get_serializer(payment)
            response_data = dict(serializer.data)
            response_data['message'] = 'To\'lov muvaffaqiyatli qayta ishlandi'
            response_data['apartment_status'] = apartment.status
            response_data['apartment_balance'] = apartment.balance
            response_data['overdue_payments'] = apartment.get_overdue_payments()

            return Response(response_data, status=status.HTTP_200_OK)
        except ValueError as e:
            return Response({'error': f'Sana formati noto\'g\'ri: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=True, methods=['get'], permission_classes=[permissions.IsAuthenticated])
    def get_overdue_payments(self, request, pk=None):
        payment = self.get_object()
        overdue_data = payment.get_overdue_payments()
        return Response({
            'payment_id': payment.id,
            'user': payment.user.fio,
            'apartment': str(payment.apartment),
            'overdue_payments': overdue_data['overdue_payments'],
            'total_overdue': overdue_data['total_overdue']
        })

    @action(detail=True, methods=['get'], permission_classes=[permissions.IsAuthenticated])
    def download_contract(self, request, pk=None):
        payment = self.get_object()
        user = payment.user
        apartment = payment.apartment
        obj = apartment.object

        doc = DocxDocument()
        doc.add_heading(f"ДАСТЛАБКИ ШАРТНОМА № {payment.id}", 0)
        doc.add_paragraph("Куп хонадонли турар-жой биноси куриш ва сотиш тугрисида")
        doc.add_paragraph(f"« 05 » Февраль 2025 йил\tҚўқон шаҳри")
        doc.add_paragraph(
            f"Қўқон шаҳар «AXLAN HOUSE» МЧЖ номидан низомга асосан фаолият юритувчи раҳбари SODIQOV XASANJON MUXSINJONOVICH "
            f"(кейинги уринларда-«Бажарувчи» деб юритилади) бир томондан ҳамда {user.fio} (келгусида «Куп хонадонли турар-жой биносининг хонадон эгаси-Буюртмачи» деб аталади) "
            f"иккинчи томондан Ўзбекистон Республикасининг «Хужалик юритувчи субъектлар фаолиятининг шартномавий-хуқуқий базаси туғрисида»ги қонунига мувофиқ мазкур шартномани қуйидагилар туғрисида туздик."
        )

        doc.add_heading("ШАРТНОМА ПРЕДМЕТИ.", level=1)
        doc.add_paragraph(
            f"1. Томонлар «Буюртмачи» хонадон сотиб олишга розилиги туғрисида «Бажарувчи» га ариза орқали мурожаат этилгандан сўнг, "
            f"Ўзбекистон Республикаси, Фарғона вилояти, Қўқон шаҳар {obj.address} да жойлашган {obj.floors} қаватли {obj.total_apartments} хонадонли "
            f"{apartment.room_number}-хонадонli турар-жой биносини қуришга, буюртмачи вазифасини бажариш тўғрисида шартномани (кейинги уринларда - асосий шартнома) тузиш мажбуриятини ўз зиммаларига оладалар."
        )

        doc.add_heading("МУҲИМ ШАРТЛАР.", level=1)
        doc.add_paragraph(
            f"а) «Буюртмачи»га топшириладиган уйнинг {apartment.room_number}-хонадон ({apartment.rooms}-хонали умумий фойдаланиш майдони {apartment.area} кв м) "
            f"умумий қийматининг бошланғич нархи {payment.total_amount} сўмни ташкил этади ва ушбу нарх томонлар томонидан келишилган ҳолда ўзгариши мумкин;"
        )
        doc.add_paragraph(
            f"б) Бажарувчи «тайёр ҳолда топшириш» шартларида турар-жой биносини қуришга бажарувчи вазифасини бажариш мажбуриятини ўз зиммасига олади..."
        )

        doc.add_heading("ХИСОБ-КИТОБ ТАРТИБИ.", level=1)
        doc.add_paragraph(
            f"«Буюртмачи» томонидан мазкур шартнома имзолангач {payment.duration_months} ой давомида яъни 31.12.2025 йилга қадар "
            f"хонадон қуришга пул ўтказиш йўли орқали банкдаги ҳисоб-варағига хонадон қийматининг 100 фоизи яъни {payment.total_amount} сўм миқдорида пул маблағини ўтказади."
        )
        if payment.payment_type == 'muddatli':
            doc.add_paragraph(
                f"Бошланғич тўлов: {payment.initial_payment} сўм, Фоиз: {payment.interest_rate}%, Ҳар ойлик тўлов: {payment.monthly_payment} сўм."
            )
        elif payment.payment_type == 'band':
            doc.add_paragraph(
                f"Band qilish uchun to‘lov: {payment.initial_payment} so‘m, Muddat: {payment.reservation_deadline}."
            )
        elif payment.payment_type == 'ipoteka':
            doc.add_paragraph(
                f"Ipoteka banki: {payment.bank_name}, Boshlang‘ich to‘lov: {payment.initial_payment} so‘m."
            )

        docx_path = os.path.join(settings.MEDIA_ROOT, f"contracts/docx/contract_{payment.id}.docx")
        os.makedirs(os.path.dirname(docx_path), exist_ok=True)
        doc.save(docx_path)

        Document.objects.create(
            payment=payment,
            docx_file=f"contracts/docx/contract_{payment.id}.docx",
            pdf_file=None
        )

        with open(docx_path, 'rb') as f:
            return FileResponse(io.BytesIO(f.read()), as_attachment=True, filename=f"contract_{payment.id}.docx")

    @action(detail=False, methods=['get'], permission_classes=[permissions.IsAuthenticated])
    def statistics(self, request):
        today = timezone.now().date()
        total_sales = Payment.objects.aggregate(total=Sum('total_amount'))['total'] or Decimal('0')
        sold_apartments = Apartment.objects.filter(status='sotilgan').count()
        clients = User.objects.filter(user_type='mijoz').count()
        total_objects = Object.objects.count()
        total_apartments = Apartment.objects.count()
        free_apartments = Apartment.objects.filter(status='bosh').count()
        reserved_apartments = Apartment.objects.filter(status='band').count()
        average_price = Apartment.objects.aggregate(avg=Avg('price'))['avg'] or Decimal('0')
        total_payments = Apartment.objects.aggregate(total=Sum('total_payments'))['total'] or Decimal('0')
        total_balance = Apartment.objects.aggregate(total=Sum('balance'))['total'] or Decimal('0')
        paid_payments = Payment.objects.filter(status='paid').aggregate(total=Sum('paid_amount'))['total'] or Decimal('0')
        pending_payments = Payment.objects.filter(status='pending').aggregate(total=Sum('monthly_payment'))['total'] or Decimal('0')
        overdue_payments = Payment.objects.filter(status='overdue').aggregate(total=Sum('monthly_payment'))['total'] or Decimal('0')
        payments_due_today = Payment.objects.filter(due_date=today.day, status='pending').aggregate(total=Sum('monthly_payment'))['total'] or Decimal('0')
        payments_paid_today = Payment.objects.filter(status='paid', created_at__date=today).aggregate(total=Sum('paid_amount'))['total'] or Decimal('0')

        data = {
            'total_sales': total_sales,
            'sold_apartments': sold_apartments,
            'clients': clients,
            'total_objects': total_objects,
            'total_apartments': total_apartments,
            'free_apartments': free_apartments,
            'reserved_apartments': reserved_apartments,
            'average_price': average_price,
            'total_payments': total_payments,
            'total_balance': total_balance,
            'paid_payments': paid_payments,
            'pending_payments': pending_payments,
            'overdue_payments': overdue_payments,
            'payments_due_today': payments_due_today,
            'payments_paid_today': payments_paid_today,
        }
        return Response(data)

class UserPaymentViewSet(viewsets.ModelViewSet):
    queryset = UserPayment.objects.all()
    serializer_class = UserPaymentSerializer
    pagination_class = CustomPagination
    permission_classes = [permissions.IsAuthenticated]
    filterset_fields = ['user', 'date', 'payment_type']

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

class SupplierPaymentViewSet(viewsets.ModelViewSet):
    queryset = SupplierPayment.objects.all()
    serializer_class = SupplierPaymentSerializer
    pagination_class = CustomPagination
    permission_classes = [permissions.IsAuthenticated]
    filterset_fields = ['supplier', 'date', 'payment_type']

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

class DocumentViewSet(viewsets.ModelViewSet):
    queryset = Document.objects.all()
    serializer_class = DocumentSerializer
    pagination_class = CustomPagination
    permission_classes = [permissions.IsAuthenticated]
    filterset_fields = ['payment', 'created_at', 'document_type']

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]


class RoomTypeModelViewSet(viewsets.ModelViewSet):
    """1x, 2x, 3x, 4x xonali ichki 3D modellar. Barcha shu xonali xonadonlar bitta modeldan foydalanadi."""
    queryset = RoomTypeModel.objects.all().order_by('room_count')
    serializer_class = RoomTypeModelSerializer
    permission_classes = [permissions.IsAuthenticated]
    pagination_class = None

    def get_permissions(self):
        if self.action in ['create', 'update', 'partial_update', 'destroy']:
            return [permissions.IsAdminUser()]
        return [permissions.IsAuthenticated()]

    @action(detail=True, methods=['post'], permission_classes=[permissions.IsAuthenticated])
    def upload_3d(self, request, pk=None):
        """Xona turi uchun ichki 3D model. .glb/.gltf yoki .zip (ichida model + teksturalar). multipart/form-data, file: model_3d"""
        rt = self.get_object()
        file = request.FILES.get('model_3d')
        if not file:
            return Response({'detail': 'model_3d fayl yuborilishi shart'}, status=status.HTTP_400_BAD_REQUEST)
        is_zip = file.name.lower().endswith('.zip')
        if not is_zip and not any(file.name.lower().endswith(ext) for ext in ('.glb', '.gltf')):
            return Response({'detail': 'Faqat .glb, .gltf yoki .zip qabul qilinadi'}, status=status.HTTP_400_BAD_REQUEST)
        target_dir = os.path.join(settings.MEDIA_ROOT, f'room_type_3d/room_{rt.room_count}')
        if rt.model_3d:
            rt.model_3d.delete(save=False)
        if os.path.isdir(target_dir):
            shutil.rmtree(target_dir, ignore_errors=True)
        if is_zip:
            main_relative = _extract_zip_and_find_model(file, target_dir)
            if not main_relative:
                return Response({'detail': 'ZIP ichida .glb yoki .gltf topilmadi yoki xatolik'}, status=status.HTTP_400_BAD_REQUEST)
            prefix = f'room_type_3d/room_{rt.room_count}/'
            suffix = main_relative.split(prefix, 1)[-1] if prefix in main_relative else os.path.basename(main_relative)
            main_path = os.path.join(settings.MEDIA_ROOT, main_relative)
            with open(main_path, 'rb') as fh:
                rt.model_3d.save(suffix, File(fh), save=True)
        else:
            rt.model_3d = file
            rt.save()
        try:
            if rt.model_3d and rt.model_3d.name:
                full_path = os.path.join(settings.MEDIA_ROOT, rt.model_3d.name)
                if os.path.isfile(full_path):
                    _convert_gltf_to_metalrough(full_path)
        except Exception:
            pass
        serializer = self.get_serializer(rt)
        return Response(serializer.data, status=status.HTTP_200_OK)

    @action(detail=True, methods=['post'], permission_classes=[permissions.IsAuthenticated])
    def delete_3d(self, request, pk=None):
        """Xona turi 3D modelini o'chirish."""
        rt = self.get_object()
        if rt.model_3d:
            rt.model_3d.delete(save=False)
            rt.model_3d = None
            rt.save()
        target_dir = os.path.join(settings.MEDIA_ROOT, f'room_type_3d/room_{rt.room_count}')
        if os.path.isdir(target_dir):
            shutil.rmtree(target_dir, ignore_errors=True)
        serializer = self.get_serializer(rt)
        return Response(serializer.data, status=status.HTTP_200_OK)


class ReportViewSet(viewsets.GenericViewSet):
    """Hisobot yuklab olish: joriy holatni generatsiya qilish va yuklab olingan hisobotlar ro'yxati."""
    queryset = OrganizationReport.objects.all().order_by('-created_at')
    serializer_class = OrganizationReportSerializer
    permission_classes = [permissions.IsAuthenticated]

    def list(self, request, *args, **kwargs):
        """Yuklab olingan hisobotlar ro'yxati."""
        queryset = self.get_queryset()
        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)

    @action(detail=False, methods=['post'], url_path='generate-full')
    def generate_full(self, request):
        """Joriy holat bo'yicha to'liq hisobot yaratadi (DOCX), saqlaydi va qaytaradi."""
        try:
            report = build_full_report_docx()
            serializer = self.get_serializer(report)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        except Exception as e:
            return Response({'detail': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=True, methods=['get'], url_path='download')
    def download(self, request, pk=None):
        """Hisobot faylini yuklab olish."""
        report = self.get_object()
        if not report.file:
            return Response({'detail': 'Fayl topilmadi.'}, status=status.HTTP_404_NOT_FOUND)
        file_path = os.path.join(settings.MEDIA_ROOT, report.file.name)
        if not os.path.isfile(file_path):
            return Response({'detail': 'Fayl diskda topilmadi.'}, status=status.HTTP_404_NOT_FOUND)
        filename = os.path.basename(report.file.name)
        with open(file_path, 'rb') as f:
            return FileResponse(io.BytesIO(f.read()), as_attachment=True, filename=filename)
